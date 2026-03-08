import os
import re
import time
import hashlib
import logging
from pathlib import Path
from typing import List
from datetime import datetime

import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

import config_data as config
from vector_stores import VectorStoreService, DashScopeEmbeddings

logger = logging.getLogger("KnowledgeBase")

def ultra_clean(text: str) -> str:
    """深度清洗文本，只保留中文/英文/数字/常用标点"""
    if not isinstance(text, str) or pd.isna(text):
        return ""
    text = "".join(c for c in text if c.isprintable())
    text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9，。！？；：（）()《》、\s\-\.]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def read_file_content(file_path: str) -> str:
    """读取 PDF / Word / Text 文件内容"""
    suffix = Path(file_path).suffix.lower()
    content = ""
    try:
        if suffix == ".pdf":
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
        elif suffix in [".docx", ".doc"]:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        elif suffix in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            logger.warning(f"⚠️ 不支持的文件类型: {file_path}")
    except Exception as e:
        logger.error(f"❌ 读取文件 {file_path} 出错: {e}")
    return content.strip()

class KnowledgeBaseService(object):
    def __init__(self):
        self.embedding = DashScopeEmbeddings(model_name=config.embedding_model_name)
        self.vector_service = VectorStoreService(embedding=self.embedding)
        
        self.spliter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            separators=config.separators,
            length_function=len,
        )

    def ingest_files(self, file_paths: List[str]):
        """
        批量索引文件到 FAISS
        """
        all_docs = []
        for path in file_paths:
            raw_text = read_file_content(path)
            clean_text = ultra_clean(raw_text)
            
            if clean_text:
                # 如果文本太长，切分它
                chunks = self.spliter.split_text(clean_text)
                for chunk in chunks:
                    all_docs.append(Document(
                        page_content=chunk,
                        metadata={
                            "source": Path(path).name,
                            "timestamp": time.time(),
                            "create_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                    ))
        
        if not all_docs:
            logger.warning("⚠️ 没有有效文档可以处理")
            return "没有有效文档"

        # 添加到向量库
        self.vector_service.vector_store.add_documents(all_docs)
        self.vector_service.save_local()
        
        logger.info(f"✅ 成功索引 {len(all_docs)} 条文本段并保存库。")
        return f"成功索引 {len(all_docs)} 条片段"

if __name__ == '__main__':
    service = KnowledgeBaseService()
    # 示例
    # res = service.ingest_files(["example.pdf"])
    # print(res)
